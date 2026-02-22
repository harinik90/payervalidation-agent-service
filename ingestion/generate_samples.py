"""
ingestion/generate_samples.py

Generate synthetic sample data for development and testing.

Creates:
    data/raw/policies/   -- DOCX coverage policy documents  (ChromaDB ingestion)
    data/raw/cases/      -- JSON prior authorization cases   (agent input fixtures)

Case pool (16 total, workflow cases always generated first):
    Workflow cases  (WF1-WF6) -- map directly to WORKFLOWS.md scenarios
    Clinical cases  (C01-C10) -- broader clinical coverage, all happy-path

All data is fully synthetic.  No real PHI.
NPI 1033472386 used in WF4 is drawn from the real OIG LEIE database (excl. type 1128a1).

Usage:
    python ingestion/generate_samples.py                  # all 16 cases + all 11 policies
    python ingestion/generate_samples.py --count 6        # first 6 cases (workflow only)
    python ingestion/generate_samples.py --count 10       # first 10 cases
    python ingestion/generate_samples.py --type workflow  # only WF1-WF6 cases
    python ingestion/generate_samples.py --type clinical  # only C01-C10 cases
    python ingestion/generate_samples.py --no-policies    # skip policy DOCX generation
    python ingestion/generate_samples.py --no-cases       # skip case JSON generation
"""
from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

import docx

DATA_DIR = Path(os.environ.get("DATA_DIR", "./data"))
POLICY_DIR = DATA_DIR / "raw" / "policies"
CASE_DIR = DATA_DIR / "raw" / "cases"


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _heading(doc: docx.Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: docx.Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullet(doc: docx.Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _save_policy(doc: docx.Document, name: str) -> None:
    POLICY_DIR.mkdir(parents=True, exist_ok=True)
    path = POLICY_DIR / name
    doc.save(str(path))
    print(f"  [policy] {path.name}")


def _save_case(data: dict, name: str) -> None:
    CASE_DIR.mkdir(parents=True, exist_ok=True)
    path = CASE_DIR / name
    path.write_text(json.dumps(data, indent=2))
    print(f"  [case  ] {path.name}")


# ── Policy documents (11 total) ───────────────────────────────────────────────

def policy_cgm_type2_diabetes() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: CONTINUOUS GLUCOSE MONITORING — TYPE 2 DIABETES")
    _para(d, "Policy Number: CP-DM-001  |  Effective Date: 2024-01-01  |  Line of Business: Commercial")
    _para(d, "ICD-10 Codes: E11, E1100, E1101, E119, E1165, E11649  |  CPT Codes: 95250, 95251, K0553, K0554")
    _para(d, "Policy Type: policy")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy establishes coverage criteria for continuous glucose monitoring (CGM) devices for members with Type 2 diabetes mellitus (T2DM). CGM provides real-time interstitial glucose readings and trend data, enabling tighter glycemic control and reducing hypoglycemic events.")

    _heading(d, "COVERAGE CRITERIA", 2)
    _para(d, "Coverage is approved when ALL of the following criteria are met:")
    _bullet(d, [
        "Member has a confirmed diagnosis of Type 2 diabetes mellitus (ICD-10: E11.xx).",
        "Member is currently using insulin therapy (basal, bolus, or combination) and requires frequent dose adjustments.",
        "HbA1c is >= 7.5% documented within the past 3 months, OR member has experienced >= 2 documented hypoglycemic episodes (blood glucose < 70 mg/dL) in the past 6 months.",
        "Member is not already covered for a CGM device through another payer or government program.",
        "Prescribing provider is a board-certified endocrinologist or primary care physician with documented diabetes management experience.",
        "Member has received diabetes self-management education (DSME) within the past 12 months.",
    ])

    _heading(d, "NON-COVERED INDICATIONS", 2)
    _bullet(d, [
        "Type 2 diabetes managed by diet and/or non-insulin oral agents only, without insulin therapy.",
        "Duplicate CGM device coverage when a functional device is already covered.",
        "CGM solely for weight management or non-diabetic glucose monitoring.",
        "Experimental CGM devices without FDA clearance.",
    ])

    _heading(d, "MEDICARE ADVANTAGE — CGM COVERAGE", 2)
    _para(d, "For Medicare Advantage members, CGM coverage follows CMS National Coverage Determination (NCD) 280.1. As of January 2025 (NCD 280.1 amendment), CGM is covered for ALL Medicare beneficiaries with diabetes (Type 1 or Type 2, insulin-dependent or non-insulin) when ordered by the treating endocrinologist or primary care physician. This supersedes prior restrictions limiting CGM to insulin-dependent patients.")

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Letter of medical necessity from treating physician.",
        "Recent HbA1c lab result (dated within 90 days).",
        "Current medication list confirming insulin therapy (commercial only).",
        "Diabetes education records or attestation.",
    ])

    _heading(d, "AUTHORIZATION PERIOD", 2)
    _para(d, "Initial authorization: 12 months. Renewal requires updated HbA1c and clinical documentation demonstrating continued use and clinical benefit.")
    _save_policy(d, "CP-DM-001_cgm_type2_diabetes.docx")


def policy_bariatric_surgery() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: BARIATRIC SURGERY")
    _para(d, "Policy Number: CP-BS-002  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicaid")
    _para(d, "ICD-10 Codes: E6601, E6609, E669, Z6841, Z6842, Z6843, Z6844, Z6845  |  CPT Codes: 43644, 43645, 43770, 43771, 43772, 43773, 43774, 43775, 43842, 43843, 43845, 43846, 43847, 43848")
    _para(d, "Policy Type: policy")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy defines coverage criteria for bariatric surgical procedures including Roux-en-Y gastric bypass (RYGB), sleeve gastrectomy, adjustable gastric banding, and biliopancreatic diversion with duodenal switch. Bariatric surgery is indicated for eligible patients who have failed conservative weight management strategies.")

    _heading(d, "COVERAGE CRITERIA", 2)
    _bullet(d, [
        "BMI >= 40 kg/m2 (morbid obesity, ICD-10 Z68.41-Z68.45), OR",
        "BMI >= 35 kg/m2 with at least one obesity-related comorbidity: Type 2 diabetes (E11), hypertension (I10), obstructive sleep apnea (G4733), osteoarthritis (M19.9), or GERD (K219).",
        "Member has participated in a medically supervised weight management program for a minimum of 6 consecutive months within the past 2 years.",
        "Psychological evaluation completed by a licensed mental health professional within the past 12 months.",
        "Medical clearance from the primary care physician.",
        "Nutritional counseling documentation from a registered dietitian (minimum 3 sessions).",
        "Member is age 18-65.",
        "No active substance use disorder within the past 12 months.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Operative note from bariatric surgeon.",
        "6-month medically supervised diet program records.",
        "Psychological clearance letter.",
        "Dietitian consultation notes.",
        "PCP medical clearance.",
        "Current height/weight/BMI from office visit within 60 days.",
    ])
    _save_policy(d, "CP-BS-002_bariatric_surgery.docx")


def policy_biologics_ra() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: BIOLOGIC THERAPY — RHEUMATOID ARTHRITIS")
    _para(d, "Policy Number: CP-RA-003  |  Effective Date: 2024-04-01  |  Line of Business: Commercial")
    _para(d, "ICD-10 Codes: M0500, M0510, M0520, M0530, M0560, M0570, M0580, M0590, M069, M0800  |  CPT/HCPCS: J0135, J0171, J0717, J1602, J3262, J2182, J0129, J0223")
    _para(d, "Policy Type: policy")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy governs prior authorization of biologic disease-modifying antirheumatic drugs (bDMARDs) for rheumatoid arthritis (RA). Covered agents include TNF inhibitors (adalimumab, etanercept, infliximab, certolizumab, golimumab), IL-6 inhibitors (tocilizumab, sarilumab), and JAK inhibitors (tofacitinib, baricitinib, upadacitinib).")

    _heading(d, "COVERAGE CRITERIA — INITIAL AUTHORIZATION", 2)
    _bullet(d, [
        "Confirmed diagnosis of moderate-to-severe rheumatoid arthritis by a board-certified rheumatologist.",
        "Disease Activity Score (DAS28) >= 3.2 documented within the past 3 months.",
        "Step therapy: member must have trialed and had inadequate response to at least TWO conventional DMARDs (methotrexate, hydroxychloroquine, sulfasalazine, leflunomide) at adequate doses for >= 3 months each.",
        "Negative TB screening (IGRA or TST) within the past 12 months.",
        "Hepatitis B surface antigen and antibody panel reviewed.",
    ])

    _heading(d, "PREFERRED AGENTS", 2)
    _bullet(d, [
        "Tier 1 (preferred): Adalimumab biosimilar (Hadlima, Hyrimoz, Cyltezo), Etanercept biosimilar.",
        "Tier 2 (require Tier 1 failure): Reference adalimumab (Humira), Certolizumab, Golimumab.",
        "Tier 3 (require Tier 1 and Tier 2 failure): IL-6 inhibitors, JAK inhibitors.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Rheumatologist office note with current DAS28 score.",
        "Prior DMARD trial history with dates, doses, and reason for discontinuation.",
        "TB test result (dated within 12 months).",
        "Hepatitis B panel results.",
    ])
    _save_policy(d, "CP-RA-003_biologics_rheumatoid_arthritis.docx")


def policy_spinal_cord_stimulation() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: SPINAL CORD STIMULATION")
    _para(d, "Policy Number: CP-PM-004  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicare Advantage")
    _para(d, "ICD-10 Codes: M5416, M5417, M54400, M54401, M54402, G8929, G89211, G89218, G8922  |  CPT Codes: 63650, 63655, 63661, 63662, 63663, 63664, 63685, 63688, 95970, 95971, 95972")
    _para(d, "Policy Type: policy")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy establishes coverage for spinal cord stimulation (SCS) as treatment for chronic intractable pain that has not responded to conventional therapies.")

    _heading(d, "COVERED INDICATIONS", 2)
    _bullet(d, [
        "Failed back surgery syndrome (FBSS) with persistent radiculopathy following lumbar surgery.",
        "Complex regional pain syndrome (CRPS) Type I or II.",
        "Refractory chronic low back pain with radicular component (duration >= 6 months).",
    ])

    _heading(d, "COVERAGE CRITERIA — TRIAL IMPLANT (CPT 63650)", 2)
    _bullet(d, [
        "Chronic pain >= 6 months duration unresponsive to conservative treatment.",
        "Failed adequate trial (>= 6 weeks each) of at least THREE conservative therapies: NSAIDs, opioids, physical therapy, TENS, nerve blocks, or epidural steroid injections.",
        "Psychological evaluation by a licensed psychologist or psychiatrist confirming suitability for implanted device.",
        "Pain management specialist evaluation confirming surgical candidacy.",
        "Trial stimulation conducted for minimum 3-7 days before permanent implant authorization.",
    ])

    _heading(d, "COVERAGE CRITERIA — PERMANENT IMPLANT (CPT 63685)", 2)
    _bullet(d, [
        "Successful trial demonstrating >= 50% reduction in pain intensity (VAS or NRS score) documented by the implanting physician.",
        "Member requests permanent implant.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Pain management or neurosurgery office notes with pain scores.",
        "Documentation of all prior conservative therapy trials (minimum 6 months).",
        "Psychological evaluation letter.",
        "Trial stimulation procedure note with documented pain reduction outcome.",
    ])
    _save_policy(d, "CP-PM-004_spinal_cord_stimulation.docx")


def policy_home_health() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: HOME HEALTH SERVICES")
    _para(d, "Policy Number: CP-HH-005  |  Effective Date: 2024-01-01  |  Line of Business: Medicare Advantage, Commercial")
    _para(d, "ICD-10 Codes: Z741, Z742, Z8901, Z9911  |  CPT Codes: 99500, 99501, 99502, 99503, 99504, 99505, 99506, 99507, 99509, 99510, 99511, 99512")
    _para(d, "Policy Type: benefit")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy defines coverage for skilled home health services provided by a Medicare-certified or state-licensed home health agency.")

    _heading(d, "COVERED SERVICES", 2)
    _bullet(d, [
        "Skilled nursing visits (wound care, IV therapy, catheter care, medication management).",
        "Physical therapy (gait training, strengthening, transfer training).",
        "Occupational therapy (ADL training, adaptive equipment, home safety evaluation).",
        "Speech-language pathology.",
        "Medical social work.",
        "Home health aide services (only when concurrent with skilled services).",
    ])

    _heading(d, "COVERAGE CRITERIA", 2)
    _bullet(d, [
        "Member is homebound: leaving home requires considerable effort and is medically inadvisable.",
        "Services are medically necessary and skilled in nature.",
        "Care is ordered by a physician or NP following a face-to-face encounter within the past 90 days.",
        "A certified home health agency develops and follows a Plan of Care (POC) signed by the ordering provider.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Physician order for home health with diagnosis and services ordered.",
        "Face-to-face encounter note within 90 days.",
        "Completed Plan of Care (Form 485 or equivalent).",
        "Documentation supporting homebound status.",
    ])
    _save_policy(d, "CP-HH-005_home_health_services.docx")


def policy_mental_health_parity() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: MENTAL HEALTH AND SUBSTANCE USE DISORDER")
    _para(d, "Policy Number: CP-BH-006  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicaid, Medicare Advantage")
    _para(d, "ICD-10 Codes: F20, F32, F33, F41, F42, F43, F50, F60, F70, F90, F10, F11, F12, F13, F14, F15, F16  |  CPT Codes: 90791, 90792, 90832, 90834, 90837, 90839, 90840, 90847, 90853, H0001, H0004, H0015, H0020")
    _para(d, "Policy Type: benefit")

    _heading(d, "PURPOSE AND REGULATORY BASIS", 2)
    _para(d, "This policy governs mental health (MH) and substance use disorder (SUD) benefits in compliance with the Mental Health Parity and Addiction Equity Act (MHPAEA) and applicable state parity laws.")

    _heading(d, "COVERED SERVICES", 2)
    _bullet(d, [
        "Outpatient individual psychotherapy: covered at same cost-sharing as primary care visit.",
        "Intensive Outpatient Program (IOP): minimum 9 hours/week, covered up to 30 days per episode.",
        "Partial Hospitalization Program (PHP): minimum 20 hours/week, covered up to 30 days per episode.",
        "Inpatient psychiatric hospitalization: acute stabilization; authorization required after 72 hours.",
        "Medication-assisted treatment (MAT): buprenorphine, naltrexone, methadone.",
        "Residential SUD treatment: authorized per ASAM placement criteria.",
    ])

    _heading(d, "PRIOR AUTHORIZATION REQUIREMENTS", 2)
    _bullet(d, [
        "Outpatient therapy (< 26 sessions/year): NO prior authorization required.",
        "IOP and PHP: authorization required with LOCUS or ASAM placement criteria documentation.",
        "Inpatient psychiatric: concurrent review required after 72 hours.",
        "Residential SUD: authorization required prior to admission.",
    ])
    _save_policy(d, "CP-BH-006_mental_health_parity.docx")


def policy_physical_therapy() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: PHYSICAL AND OCCUPATIONAL THERAPY")
    _para(d, "Policy Number: CP-PT-007  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicare Advantage")
    _para(d, "ICD-10 Codes: M1711, M1712, M1721, M1722, M5416, M5417, M2551, M2552, S72001A, S72002A, M1611, M1612  |  CPT Codes: 97110, 97112, 97116, 97140, 97150, 97530, 97535, 97542")
    _para(d, "Policy Type: benefit")

    _heading(d, "COVERAGE CRITERIA", 2)
    _bullet(d, [
        "Services are prescribed by a licensed physician, NP, PA, or specialist.",
        "Conditions are expected to improve with skilled therapy intervention.",
        "Services require the skills of a licensed PT/OT.",
        "Evaluation completed within 30 days of referral with documented functional baseline.",
    ])

    _heading(d, "AUTHORIZED VISIT LIMITS", 2)
    _bullet(d, [
        "Initial authorization: up to 12 visits per condition per 90-day period.",
        "Extended authorization (visits 13-30): requires functional progress note showing >= 25% improvement.",
        "Maximum annual benefit: 60 visits PT + 60 visits OT.",
        "Post-surgical rehabilitation: additional 30 visits authorized for major joint replacement.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION FOR EXTENDED AUTHORIZATION", 2)
    _bullet(d, [
        "Initial evaluation with baseline functional scores.",
        "Progress note documenting measurable functional improvement.",
        "Updated treatment plan with specific, time-limited goals.",
    ])
    _save_policy(d, "CP-PT-007_physical_occupational_therapy.docx")


def policy_insulin_pump() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: INSULIN INFUSION PUMP — DURABLE MEDICAL EQUIPMENT")
    _para(d, "Policy Number: CP-DM-008  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicare Advantage")
    _para(d, "ICD-10 Codes: E1010, E1011, E1040, E1041, E1065, E10649  |  HCPCS Codes: E0784, A9274, A9276, A9277, A9278, K0553, K0554")
    _para(d, "Policy Type: policy")

    _heading(d, "COVERAGE CRITERIA", 2)
    _bullet(d, [
        "Confirmed diagnosis of Type 1 diabetes mellitus (ICD-10: E10.xx), OR Type 2 diabetes requiring intensive insulin therapy.",
        "Member currently managed with multiple daily injection (MDI) therapy using both basal and bolus insulin.",
        "HbA1c >= 7.0% despite MDI therapy, OR documented recurrent severe hypoglycemia, OR hypoglycemia unawareness.",
        "Prescribing provider is a board-certified endocrinologist.",
        "Member has completed formal pump training (minimum 4 hours) with a certified diabetes educator.",
        "Member agrees to perform >= 4 daily fingerstick glucose checks, OR uses an approved integrated CGM.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Endocrinologist letter of medical necessity.",
        "HbA1c results for past 6 months.",
        "Current insulin regimen documentation.",
        "Pump training completion certificate from CDE.",
        "Self-monitoring blood glucose logs (past 30 days).",
    ])
    _save_policy(d, "CP-DM-008_insulin_pump_dme.docx")


def policy_colonoscopy() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: COLONOSCOPY AND COLORECTAL CANCER SCREENING")
    _para(d, "Policy Number: CP-GI-009  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicare Advantage, Medicaid")
    _para(d, "ICD-10 Codes: Z1211, Z1212, K631, K638, C180, C181, C182, C183, C184, C185, C186, C187, C188, C189, D1200, D1201  |  CPT Codes: 45378, 45380, 45381, 45382, 45384, 45385, 45386, 45388, 45391, 45392, G0105, G0121")
    _para(d, "Policy Type: benefit")

    _heading(d, "PREVENTIVE SCREENING — NO COST SHARING", 2)
    _bullet(d, [
        "Screening colonoscopy every 10 years beginning at age 45 for average-risk members.",
        "High-risk screening (first-degree relative with CRC before age 60): every 5 years beginning at age 40.",
        "Follow-up colonoscopy after a positive FIT or Cologuard test: treated as preventive at 100% coverage.",
    ])

    _heading(d, "PRIOR AUTHORIZATION", 2)
    _para(d, "No prior authorization required for screening colonoscopy. Authorization required for:")
    _bullet(d, [
        "Therapeutic polypectomy or endoscopic resection procedures (45385, 45386, 45388).",
        "Repeat colonoscopy within 12 months of a previous procedure.",
        "Surveillance intervals shorter than guideline-recommended.",
    ])
    _save_policy(d, "CP-GI-009_colonoscopy_crc_screening.docx")


def policy_oncology_immunotherapy() -> None:
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: ONCOLOGY — PD-1/PD-L1 CHECKPOINT INHIBITOR IMMUNOTHERAPY")
    _para(d, "Policy Number: CP-ONC-010  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicare Advantage")
    _para(d, "ICD-10 Codes: C3490, C3491, C3492, C3410, C3411, C3412, C349, C780, C781, C782, C7800  |  HCPCS Codes: J9271, J9299, J9177, J0129, J2357")
    _para(d, "Policy Type: policy")

    _heading(d, "COVERAGE CRITERIA", 2)
    _bullet(d, [
        "FDA-approved indication OR NCCN Category 1 or 2A recommendation for specific tumor type, stage, and biomarker status.",
        "Required biomarker testing completed prior to authorization: PD-L1 IHC, MSI/MMR testing, TMB where applicable.",
        "ECOG performance status <= 2 documented within 30 days.",
        "No active autoimmune disease requiring systemic immunosuppression.",
        "Prescribing oncologist is board-certified.",
    ])

    _heading(d, "COVERED INDICATIONS — PEMBROLIZUMAB (J9271)", 2)
    _bullet(d, [
        "NSCLC: first-line monotherapy for PD-L1 TPS >= 50%, no EGFR/ALK alterations.",
        "NSCLC combination with platinum chemotherapy for PD-L1 >= 1%.",
        "Triple-negative breast cancer: combination with chemotherapy for PD-L1 CPS >= 10.",
        "MSI-H/dMMR solid tumors: tissue-agnostic second-line.",
    ])

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Pathology report confirming tumor type, stage, and biomarker results.",
        "Oncologist office note with current regimen and ECOG PS.",
        "For renewal: most recent imaging report and response assessment.",
    ])
    _save_policy(d, "CP-ONC-010_checkpoint_inhibitor_immunotherapy.docx")


def policy_acupuncture_exclusion() -> None:
    """WF3 support: plan exclusion document for non-covered complementary services."""
    d = docx.Document()
    _heading(d, "BENEFIT POLICY: NON-COVERED COMPLEMENTARY AND ALTERNATIVE MEDICINE SERVICES")
    _para(d, "Policy Number: CP-COMP-011  |  Effective Date: 2024-01-01  |  Line of Business: Commercial")
    _para(d, "ICD-10 Codes: M545, M5450, M5459, G893, G8929  |  CPT Codes: 97810, 97811, 97813, 97814, 97140, 97533, S8948")
    _para(d, "Policy Type: benefit")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy identifies complementary and alternative medicine (CAM) services that are explicitly excluded from coverage under standard commercial HMO and PPO benefit plans. Excluded services are not reimbursable regardless of medical necessity documentation unless a specific rider or supplemental benefit has been purchased.")

    _heading(d, "NON-COVERED SERVICES", 2)
    _para(d, "The following services are explicitly excluded from coverage under this commercial benefit plan:")
    _bullet(d, [
        "Acupuncture (CPT 97810, 97811, 97813, 97814) -- all indications including chronic pain, nausea, and headache.",
        "Acupressure and traditional Chinese medicine modalities.",
        "Massage therapy for wellness or chronic pain (CPT S8948).",
        "Chiropractic maintenance care beyond 12 visits/year (acute, medically necessary chiropractic is covered separately).",
        "Naturopathic medicine services.",
        "Homeopathy, herbal medicine, and nutritional supplement counseling.",
        "Biofeedback (CPT 97533) for stress management (medically necessary biofeedback for urinary incontinence covered under CP-URO-015).",
        "Energy healing modalities (Reiki, therapeutic touch, magnet therapy).",
    ])

    _heading(d, "REGULATORY NOTE", 2)
    _para(d, "As of the effective date of this policy, no federal mandate (ACA, ERISA) or Florida state insurance mandate (Chapter 627, F.S.) requires commercial health plans to cover acupuncture services. Medicare Advantage plans may cover acupuncture for chronic low back pain under a supplemental benefit; this exclusion applies to commercial plans only. If a future state or federal mandate requires coverage, this policy will be updated and the mandate takes precedence.")

    _heading(d, "MEMBER RIGHTS", 2)
    _para(d, "Members may appeal a denial based on this exclusion. However, an exclusion from the certificate of coverage is not subject to the clinical criteria appeal process. Members wishing to purchase acupuncture coverage may inquire about available supplemental benefit riders during open enrollment.")
    _save_policy(d, "CP-COMP-011_cam_exclusions_commercial.docx")


def policy_ortho_knee_arthroplasty() -> None:
    """WF1 and WF6 support: total knee arthroplasty policy."""
    d = docx.Document()
    _heading(d, "COVERAGE POLICY: TOTAL KNEE ARTHROPLASTY (TOTAL KNEE REPLACEMENT)")
    _para(d, "Policy Number: CP-ORTHO-012  |  Effective Date: 2024-01-01  |  Line of Business: Commercial, Medicare Advantage")
    _para(d, "ICD-10 Codes: M1711, M1712, M1721, M1722, M1731, M1732, M179  |  CPT Codes: 27447, 27446, 27487, 27486")
    _para(d, "Policy Type: policy")

    _heading(d, "PURPOSE", 2)
    _para(d, "This policy establishes prior authorization criteria for total knee arthroplasty (TKA, CPT 27447) and related knee replacement procedures. TKA is indicated for end-stage knee osteoarthritis or other qualifying knee pathology when conservative treatment has been exhausted.")

    _heading(d, "COVERAGE CRITERIA", 2)
    _para(d, "ALL of the following criteria must be met:")
    _bullet(d, [
        "Diagnosis of primary osteoarthritis of the knee (ICD-10 M17.11, M17.12, M17.21, M17.22) or post-traumatic arthritis, confirmed by radiographic imaging.",
        "Radiographic evidence of Grade III or IV Kellgren-Lawrence osteoarthritis (joint space narrowing, osteophytes, subchondral sclerosis).",
        "Failed conservative therapy: member has attempted and documented inadequate response to ALL of the following for >= 6 months: physical therapy (>= 8 weeks), weight management counseling or supervised weight loss program (if BMI >= 30), oral analgesics (NSAIDs or acetaminophen at adequate dosing), and at least ONE of: intra-articular corticosteroid injections, hyaluronic acid injections, or bracing.",
        "Functional impairment documented: significant limitation in activities of daily living, ambulation, or work-related activities despite conservative treatment.",
        "Member is medically appropriate for general or regional anesthesia: cardiac, pulmonary, and coagulation clearance from PCP or specialist.",
        "BMI <= 45 kg/m2 (BMI > 45 requires additional medical optimization documentation).",
    ])

    _heading(d, "NON-COVERED", 2)
    _bullet(d, [
        "TKA for isolated anterior knee pain without structural joint disease.",
        "Bilateral simultaneous TKA without specific clinical justification (staged bilateral covered; simultaneous requires additional review).",
        "Unicompartmental knee arthroplasty as a primary procedure when total knee is indicated per imaging.",
    ])

    _heading(d, "CODING GUIDANCE", 2)
    _para(d, "CPT 27447 (total knee arthroplasty) includes all components of the arthroplasty procedure. Per CCI edits, the following codes are bundled into 27447 and cannot be billed separately on the same date of service: 27370 (injection of knee joint), 27310 (arthrotomy, knee), 27330 (arthrotomy, knee for synovial biopsy). Submission of these codes alongside 27447 will result in the claim being returned for coding correction.")

    _heading(d, "REQUIRED DOCUMENTATION", 2)
    _bullet(d, [
        "Orthopedic surgeon operative note with planned procedure.",
        "Knee X-rays (weight-bearing, minimum 2 views) with radiologist or surgeon interpretation confirming Grade III/IV OA.",
        "Conservative therapy records demonstrating >= 6 months of failed treatment.",
        "Medical clearance from PCP.",
        "BMI documentation from office visit within 60 days.",
    ])
    _save_policy(d, "CP-ORTHO-012_total_knee_arthroplasty.docx")


# ── Case pool definition ──────────────────────────────────────────────────────

# Cases tagged with "workflow" are the 6 WORKFLOWS.md scenarios.
# Cases tagged "clinical" are the broader happy-path clinical set.

_WORKFLOW_CASES: list[dict] = [
    # ── WF1: Standard APPROVE — right TKA, full documentation ────────────────
    {
        "case_id": "WF-001",
        "workflow": "WF1_APPROVE",
        "expected_decision": "APPROVE",
        "member_id": "M-1234567",
        "member_name": "Harold R. Cooper",
        "dob": "1958-06-14",
        "lob": "commercial",
        "npi": "1234567890",
        "provider_name": "Riverside Orthopedic Group",
        "facility_npi": "9876540001",
        "facility_name": "Riverside Orthopedic Surgery Center",
        "service_date": "2025-10-15",
        "icd10_codes": ["M1711"],
        "cpt_codes": ["27447"],
        "clinical_notes": (
            "Patient is a 66-year-old male with right knee primary osteoarthritis confirmed on "
            "weight-bearing AP and lateral X-rays demonstrating Grade IV Kellgren-Lawrence changes "
            "(complete medial joint space obliteration, large osteophytes, subchondral sclerosis). "
            "Conservative therapy documented and failed over 8 months: physical therapy (16 sessions, "
            "inadequate relief), NSAIDs (meloxicam 15mg daily x 6 months, GI tolerability issues), "
            "two intra-articular corticosteroid injections (March and June 2025, transient relief only), "
            "weight management counseling (BMI 28.4). Functional impairment: unable to walk > 1 block, "
            "difficulty with stairs, limitations in ADLs. PCP medical clearance obtained. "
            "Cardiac clearance: normal EKG, cardiologist cleared for surgery. "
            "Requesting total knee arthroplasty (CPT 27447), right knee."
        ),
        "procedure_description": "Right Total Knee Arthroplasty",
        "policy_reference": "CP-ORTHO-012",
        "state": "CA",
    },
    # ── WF2: PEND — SCS implant, documentation incomplete ────────────────────
    {
        "case_id": "WF-002",
        "workflow": "WF2_PEND_MISSING_DOCS",
        "expected_decision": "PEND",
        "member_id": "M-7654321",
        "member_name": "Diane L. Foster",
        "dob": "1967-03-29",
        "lob": "commercial",
        "npi": "2233445566",
        "provider_name": "Advanced Spine and Neurology Center",
        "facility_npi": "9876540002",
        "facility_name": "Advanced Spine Surgery Institute",
        "service_date": "2025-10-28",
        "icd10_codes": ["M54400"],
        "cpt_codes": ["63685"],
        "clinical_notes": (
            "Patient referred for spinal cord stimulator (SCS) permanent implant for chronic low back pain "
            "with bilateral lower extremity radiation. Pain duration approximately 18 months following "
            "failed conservative management. Requesting authorization for permanent SCS implant (CPT 63685). "
            "Note: Supporting documentation including conservative therapy records, psychological evaluation, "
            "and SCS trial procedure note were not included with this submission."
        ),
        "procedure_description": "Spinal Cord Stimulator Permanent Implant — Documentation Incomplete",
        "policy_reference": "CP-PM-004",
        "state": "TX",
        "submission_note": "INCOMPLETE: Missing conservative_therapy_notes_6mo, psychological_evaluation_report, scs_trial_operative_report",
    },
    # ── WF3: DENY — acupuncture, plan exclusion, FL ───────────────────────────
    {
        "case_id": "WF-003",
        "workflow": "WF3_DENY_PLAN_EXCLUSION",
        "expected_decision": "DENY",
        "member_id": "M-2223334",
        "member_name": "Gloria M. Perez",
        "dob": "1975-09-11",
        "lob": "commercial",
        "npi": "3344556677",
        "provider_name": "Wellness Integrative Health",
        "facility_npi": "9876540003",
        "facility_name": "Wellness Integrative Health Center",
        "service_date": "2025-09-10",
        "icd10_codes": ["M545"],
        "cpt_codes": ["97810"],
        "clinical_notes": (
            "Patient is a 49-year-old female requesting acupuncture treatment for chronic low back pain. "
            "Patient preference for non-pharmacologic therapy. No prior acupuncture treatment. "
            "Provider is a licensed acupuncturist (LAc). "
            "Clinical notes state: chronic low back pain, patient preference for acupuncture, "
            "patient declines pharmacologic options."
        ),
        "procedure_description": "Acupuncture — Chronic Low Back Pain",
        "policy_reference": "CP-COMP-011",
        "state": "FL",
    },
    # ── WF4: HARD STOP — sanctioned provider, NPI on OIG LEIE ────────────────
    {
        "case_id": "WF-004",
        "workflow": "WF4_HARD_STOP_OIG",
        "expected_decision": "DENIED_HARD_STOP",
        "member_id": "M-9876543",
        "member_name": "James T. Bradley",
        "dob": "1962-01-15",
        "lob": "medicaid",
        "npi": "1033472386",
        "provider_name": "Sunshine Pain Management LLC",
        "facility_npi": "9876540004",
        "facility_name": "Sunshine Pain Management Clinic",
        "service_date": "2025-11-01",
        "icd10_codes": ["M545"],
        "cpt_codes": ["64483"],
        "clinical_notes": (
            "Medicaid member requesting fluoroscopically guided lumbar epidural steroid injection "
            "for chronic radicular low back pain. ICD-10 M54.5. Provider NPI 1033472386. "
            "NOTE FOR TESTING: NPI 1033472386 is present in the OIG LEIE database "
            "(excl_type 1128a1, excl_date 20260219). This case should trigger a hard stop "
            "at the sanctions check and halt all downstream processing."
        ),
        "procedure_description": "Lumbar Epidural Steroid Injection — LEIE-Excluded Provider",
        "policy_reference": None,
        "state": "OH",
        "leie_test_note": "NPI 1033472386 confirmed in leie.db (excl_type=1128a1, excl_date=20260219)",
    },
    # ── WF5: PEND (regulatory override) — CGM, MA, non-insulin T2D ───────────
    {
        "case_id": "WF-005",
        "workflow": "WF5_PEND_REGULATORY_OVERRIDE",
        "expected_decision": "PEND",
        "member_id": "M-4445556",
        "member_name": "Evelyn N. Carter",
        "dob": "1950-04-22",
        "lob": "medicare_advantage",
        "npi": "5566778899",
        "provider_name": "Lakeside Endocrinology",
        "facility_npi": "9876540005",
        "facility_name": "Lakeside Endocrinology Associates",
        "service_date": "2025-08-05",
        "icd10_codes": ["E119"],
        "cpt_codes": ["95250"],
        "clinical_notes": (
            "Patient is a 75-year-old female Medicare Advantage member with Type 2 diabetes mellitus "
            "without complications (ICD-10 E11.9). Currently managed with metformin 1000mg BID and "
            "empagliflozin 10mg daily. Patient is NOT on insulin therapy. HbA1c: 7.8% (2025-07-15). "
            "No hypoglycemic episodes. Requesting continuous glucose monitoring (CGM, Dexcom G7, "
            "CPT 95250) for glucose trend monitoring and optimization of oral medication therapy. "
            "Ordering provider: Dr. Karen Mills, MD, board-certified endocrinologist. "
            "Note: Internal policy CP-DM-001 restricts commercial CGM to insulin-dependent patients. "
            "Medicare Advantage CGM eligibility should be evaluated against CMS NCD 280.1 (as amended "
            "January 2025), which expanded coverage to all Medicare T2D patients regardless of insulin use."
        ),
        "procedure_description": "CGM — Medicare Advantage Non-Insulin T2D (Regulatory Override Test)",
        "policy_reference": "CP-DM-001",
        "state": "OH",
        "regulatory_test_note": "Expected: policy DENY overridden by CMS NCD 280.1 amendment (eff. 2025-01-01) -> PEND for manual review",
    },
    # ── WF6: RETURNED — TKA with CCI bundle error + redundant dx ─────────────
    {
        "case_id": "WF-006",
        "workflow": "WF6_RETURNED_CODING_ERROR",
        "expected_decision": "RETURNED_FOR_CORRECTION",
        "member_id": "M-3334455",
        "member_name": "Arthur P. Quinn",
        "dob": "1955-11-30",
        "lob": "commercial",
        "npi": "4455667788",
        "provider_name": "Metro Orthopaedic Surgery",
        "facility_npi": "9876540006",
        "facility_name": "Metro Surgical Center",
        "service_date": "2025-11-12",
        "icd10_codes": ["M1711", "M25361"],
        "cpt_codes": ["27447", "27370"],
        "clinical_notes": (
            "Patient is a 69-year-old male with right knee primary osteoarthritis (M17.11), Grade IV "
            "on weight-bearing X-ray. Failed conservative therapy. Requesting right total knee arthroplasty "
            "(CPT 27447) with intraoperative knee injection (CPT 27370). Secondary diagnosis: stiffness of "
            "right knee (M25.361). "
            "NOTE FOR TESTING: This submission contains two intentional coding errors: "
            "(1) CPT 27370 is a CCI component of 27447 and cannot be billed separately on the same DOS; "
            "(2) ICD-10 M25.361 (stiffness of right knee) is clinically redundant when M17.11 is the "
            "primary diagnosis. Expected outcome: RETURNED_FOR_CORRECTION before policy review begins."
        ),
        "procedure_description": "Right TKA with Intentional CCI Bundle Error (Coding Test)",
        "policy_reference": "CP-ORTHO-012",
        "state": "IL",
        "coding_errors": [
            {"code": "27370", "type": "CCI_BUNDLE", "bundled_into": "27447", "action": "REMOVE"},
            {"code": "M25361", "type": "REDUNDANT_DX", "primary": "M1711", "action": "REVIEW"},
        ],
    },
]

_CLINICAL_CASES: list[dict] = [
    {
        "case_id": "CASE-2024-001",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100001",
        "member_name": "Jane A. Doe",
        "dob": "1972-04-15",
        "lob": "commercial",
        "npi": "1234567890",
        "provider_name": "Dr. Sarah Chen, MD — Endocrinology Associates",
        "facility_npi": "9876543210",
        "facility_name": "Riverside Medical Center",
        "service_date": "2025-03-01",
        "icd10_codes": ["E1100", "E1165"],
        "cpt_codes": ["95250", "K0553"],
        "clinical_notes": (
            "Patient is a 52-year-old female with a 14-year history of Type 2 diabetes mellitus, "
            "currently on basal-bolus insulin regimen (glargine 30 units QHS, aspart 8 units TID with meals). "
            "Most recent HbA1c: 9.2% (dated 2025-02-10). Patient has experienced 4 symptomatic hypoglycemic "
            "episodes in the past 3 months (BG < 65 mg/dL), including one episode requiring assistance. "
            "Requesting authorization for CGM (Dexcom G7). Patient completed DSME program January 2025."
        ),
        "procedure_description": "Continuous Glucose Monitoring — Dexcom G7 CGM System",
        "policy_reference": "CP-DM-001",
        "state": "CA",
    },
    {
        "case_id": "CASE-2024-002",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100002",
        "member_name": "Robert B. Smith",
        "dob": "1978-09-22",
        "lob": "commercial",
        "npi": "1234567891",
        "provider_name": "Dr. Michael Torres, MD — Bariatric Surgery Center",
        "facility_npi": "9876543211",
        "facility_name": "Metro Surgical Institute",
        "service_date": "2025-04-15",
        "icd10_codes": ["E6601", "E1100", "I10", "Z6843"],
        "cpt_codes": ["43644"],
        "clinical_notes": (
            "Patient is a 46-year-old male with BMI 42.3 kg/m2, T2D (HbA1c 8.8%), hypertension, "
            "and obstructive sleep apnea. Completed 7-month medically supervised weight management program. "
            "Psychological evaluation cleared patient. Nutritional counseling (4 sessions) completed. "
            "PCP medical clearance obtained. Requesting Roux-en-Y gastric bypass (CPT 43644)."
        ),
        "procedure_description": "Roux-en-Y Gastric Bypass",
        "policy_reference": "CP-BS-002",
        "state": "NY",
    },
    {
        "case_id": "CASE-2024-003",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100003",
        "member_name": "Patricia C. Johnson",
        "dob": "1965-11-08",
        "lob": "commercial",
        "npi": "1234567892",
        "provider_name": "Dr. Aisha Williams, MD — Rheumatology Specialists",
        "facility_npi": "9876543212",
        "facility_name": "University Rheumatology Clinic",
        "service_date": "2025-02-20",
        "icd10_codes": ["M0510", "M0520"],
        "cpt_codes": ["J0135"],
        "clinical_notes": (
            "Patient is a 59-year-old female with seropositive RA (DAS28-CRP: 4.8). "
            "Prior DMARD therapy: MTX 20mg/week x 18 months (hepatotoxicity), HCQ 400mg/day x 12 months "
            "(inadequate response). TB IGRA negative. HBsAg negative. "
            "Requesting adalimumab biosimilar (Hadlima, J0135)."
        ),
        "procedure_description": "Adalimumab Biosimilar (Hadlima) — Biologic DMARD for RA",
        "policy_reference": "CP-RA-003",
        "state": "IL",
    },
    {
        "case_id": "CASE-2024-004",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100004",
        "member_name": "Thomas D. Martinez",
        "dob": "1958-03-17",
        "lob": "medicare_advantage",
        "npi": "1234567893",
        "provider_name": "Dr. Kevin Park, MD — Neurosurgery & Pain Management",
        "facility_npi": "9876543213",
        "facility_name": "Advanced Spine Center",
        "service_date": "2025-05-01",
        "icd10_codes": ["M5416", "G89211", "M5417"],
        "cpt_codes": ["63650", "63685"],
        "clinical_notes": (
            "Patient is a 66-year-old male with FBSS following L4-L5 discectomy (2021) and L3-S1 fusion (2022). "
            "Persistent bilateral radiculopathy, NRS pain 8/10. Conservative therapy failed: NSAIDs, opioids, "
            "PT (24 sessions), 3 ESI series. Psychological evaluation cleared (Dr. Amy Cho, PhD, 2025-01-05). "
            "Requesting SCS trial (CPT 63650); if >= 50% pain reduction, permanent implant to follow."
        ),
        "procedure_description": "Spinal Cord Stimulator Trial and Permanent Implant",
        "policy_reference": "CP-PM-004",
        "state": "FL",
    },
    {
        "case_id": "CASE-2024-005",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100005",
        "member_name": "Eleanor E. Thompson",
        "dob": "1940-07-30",
        "lob": "medicare_advantage",
        "npi": "1234567894",
        "provider_name": "Dr. James Wilson, MD — Orthopedics",
        "facility_npi": "9876543214",
        "facility_name": "Riverside Orthopedic Hospital",
        "service_date": "2025-03-15",
        "icd10_codes": ["Z741", "Z9911", "M1611"],
        "cpt_codes": ["99500", "99506"],
        "clinical_notes": (
            "84-year-old female, 5 days post left total hip arthroplasty. Homebound: cannot ambulate "
            "without walker. Ordering skilled nursing (wound care, warfarin monitoring), PT 3x/week, "
            "OT 2x/week. Face-to-face encounter completed day of discharge. Plan of Care signed by surgeon. "
            "Requesting 60-day home health episode."
        ),
        "procedure_description": "Home Health Services — Post-Hip Replacement Rehabilitation",
        "policy_reference": "CP-HH-005",
        "state": "MA",
    },
    {
        "case_id": "CASE-2024-006",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100006",
        "member_name": "Marcus F. Davis",
        "dob": "1990-12-05",
        "lob": "commercial",
        "npi": "1234567895",
        "provider_name": "Dr. Rebecca Stone, MD — Psychiatry",
        "facility_npi": "9876543215",
        "facility_name": "Northside Behavioral Health Center",
        "service_date": "2025-02-18",
        "icd10_codes": ["F331", "F41", "F1010"],
        "cpt_codes": ["90837", "90853", "H0015"],
        "clinical_notes": (
            "34-year-old male, PHQ-9: 18/27, passive SI without plan, alcohol use disorder mild (last drink 5 days ago). "
            "LOCUS score: 19 (High Intensity Community-Based). Does not meet inpatient criteria. "
            "Prior outpatient therapy: 32 sessions CBT (insufficient). "
            "Requesting PHP: 20 hours/week (individual therapy, group therapy, medication management, SUD counseling)."
        ),
        "procedure_description": "Partial Hospitalization Program (PHP) — Mental Health and SUD",
        "policy_reference": "CP-BH-006",
        "state": "GA",
    },
    {
        "case_id": "CASE-2024-007",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100007",
        "member_name": "Sandra G. Lee",
        "dob": "1968-02-14",
        "lob": "commercial",
        "npi": "1234567896",
        "provider_name": "Dr. David Brown, MD — Orthopedics",
        "facility_npi": "9876543216",
        "facility_name": "Community Orthopedic Center",
        "service_date": "2025-01-20",
        "icd10_codes": ["M1711", "M25561", "M25562"],
        "cpt_codes": ["97110", "97140", "97530"],
        "clinical_notes": (
            "57-year-old female, 3 weeks post right TKA (done 2024-12-30). "
            "Baseline LEFS: 28/80. Goals: ROM 0-120 deg, quad strength, independent ADLs. "
            "At visit 12: LEFS improved to 42/80 (50% improvement), ROM 0-105 deg. "
            "Requesting additional 18 PT visits to complete rehabilitation."
        ),
        "procedure_description": "Outpatient Physical Therapy — Post-Total Knee Arthroplasty",
        "policy_reference": "CP-PT-007",
        "state": "WA",
    },
    {
        "case_id": "CASE-2024-008",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100008",
        "member_name": "Christopher H. Nguyen",
        "dob": "1985-06-20",
        "lob": "commercial",
        "npi": "1234567897",
        "provider_name": "Dr. Jennifer Walsh, MD — Endocrinology",
        "facility_npi": "9876543217",
        "facility_name": "Diabetes & Endocrine Specialists",
        "service_date": "2025-03-10",
        "icd10_codes": ["E1040", "E1065"],
        "cpt_codes": ["E0784", "A9274", "A9276"],
        "clinical_notes": (
            "39-year-old male, T1D since age 12. MDI: glargine 22u QHS + lispro TID. "
            "HbA1c: 8.4%. Hypoglycemia unawareness (Clarke score 5), 3 severe episodes past 6 months. "
            "Pump training completed (4-hr CDE program, 2025-02-28). SMBG logs 30 days compliant >= 4/day. "
            "Requesting Omnipod 5 insulin pump (HCPCS E0784)."
        ),
        "procedure_description": "Insulin Infusion Pump (Omnipod 5) — T1D with Hypoglycemia Unawareness",
        "policy_reference": "CP-DM-008",
        "state": "TX",
    },
    {
        "case_id": "CASE-2024-009",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100009",
        "member_name": "Barbara I. Wilson",
        "dob": "1950-08-11",
        "lob": "medicare_advantage",
        "npi": "1234567898",
        "provider_name": "Dr. Steven Grant, MD — Gastroenterology",
        "facility_npi": "9876543218",
        "facility_name": "Digestive Disease Associates",
        "service_date": "2025-04-05",
        "icd10_codes": ["Z1211", "D1201", "K631"],
        "cpt_codes": ["45385"],
        "clinical_notes": (
            "74-year-old female. Prior colonoscopy (2022-09-15): 3 tubular adenomas removed (0.8 cm, 0.6 cm, 0.4 cm). "
            "Per ACG 2023 guidelines: high-risk adenoma -> 3-year surveillance interval. "
            "Requesting 3-year surveillance colonoscopy with polypectomy (CPT 45385). "
            "FH: mother with colon cancer at age 67."
        ),
        "procedure_description": "Surveillance Colonoscopy with Polypectomy — High-Risk Adenoma History",
        "policy_reference": "CP-GI-009",
        "state": "AZ",
    },
    {
        "case_id": "CASE-2024-010",
        "workflow": "clinical",
        "expected_decision": "APPROVE",
        "member_id": "MBR-100010",
        "member_name": "William J. Garcia",
        "dob": "1955-01-25",
        "lob": "commercial",
        "npi": "1234567899",
        "provider_name": "Dr. Priya Sharma, MD — Thoracic Oncology",
        "facility_npi": "9876543219",
        "facility_name": "Comprehensive Cancer Center",
        "service_date": "2025-03-25",
        "icd10_codes": ["C3492", "C7800", "Z8501"],
        "cpt_codes": ["J9271"],
        "clinical_notes": (
            "70-year-old male, Stage IV NSCLC adenocarcinoma, right lower lobe with bilateral pulmonary mets. "
            "EGFR/ALK/ROS1/KRAS neg. PD-L1 TPS 78% (22C3). TMB 12 mut/Mb. ECOG PS: 1. "
            "Requesting first-line pembrolizumab monotherapy (Keytruda, J9271) per NCCN Cat 1. "
            "No autoimmune disease. Dose: 200mg IV q3w x 8 cycles (6-month auth)."
        ),
        "procedure_description": "Pembrolizumab (Keytruda) First-Line — Stage IV NSCLC PD-L1 High",
        "policy_reference": "CP-ONC-010",
        "state": "CO",
    },
]

# Full pool: workflow cases first, then clinical cases
ALL_CASES: list[dict] = _WORKFLOW_CASES + _CLINICAL_CASES


# ── Policy generator registry ─────────────────────────────────────────────────

ALL_POLICIES = [
    policy_cgm_type2_diabetes,
    policy_bariatric_surgery,
    policy_biologics_ra,
    policy_spinal_cord_stimulation,
    policy_home_health,
    policy_mental_health_parity,
    policy_physical_therapy,
    policy_insulin_pump,
    policy_colonoscopy,
    policy_oncology_immunotherapy,
    policy_acupuncture_exclusion,     # WF3
    policy_ortho_knee_arthroplasty,   # WF1, WF6
]


# ── CLI ───────────────────────────────────────────────────────────────────────

def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Generate synthetic PA case and policy documents for payerai-gpt.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python ingestion/generate_samples.py                  # all 16 cases + all 12 policies
  python ingestion/generate_samples.py --count 6        # first 6 cases (workflow scenarios only)
  python ingestion/generate_samples.py --count 10       # first 10 cases
  python ingestion/generate_samples.py --type workflow  # only the 6 WORKFLOWS.md scenarios
  python ingestion/generate_samples.py --type clinical  # only the 10 clinical happy-path cases
  python ingestion/generate_samples.py --no-policies    # skip DOCX generation
  python ingestion/generate_samples.py --no-cases       # skip JSON generation
        """,
    )
    p.add_argument(
        "--count", "-n",
        type=int,
        default=None,
        metavar="N",
        help="Number of cases to generate from the full pool (workflow cases first). Default: all.",
    )
    p.add_argument(
        "--type",
        choices=["all", "workflow", "clinical"],
        default="all",
        help="Case subset to generate: all (default), workflow (WF1-WF6 only), clinical (C01-C10 only).",
    )
    p.add_argument(
        "--no-policies",
        action="store_true",
        help="Skip generating policy DOCX files.",
    )
    p.add_argument(
        "--no-cases",
        action="store_true",
        help="Skip generating case JSON files.",
    )
    return p.parse_args()


def _select_cases(args: argparse.Namespace) -> list[dict]:
    if args.type == "workflow":
        pool = _WORKFLOW_CASES
    elif args.type == "clinical":
        pool = _CLINICAL_CASES
    else:
        pool = ALL_CASES

    if args.count is not None:
        if args.count < 1:
            print("ERROR: --count must be >= 1", file=sys.stderr)
            sys.exit(1)
        if args.count > len(pool):
            print(
                f"WARNING: --count {args.count} exceeds pool size {len(pool)}; "
                f"generating all {len(pool)} cases.",
                file=sys.stderr,
            )
        pool = pool[: args.count]

    return pool


def main() -> None:
    args = _parse_args()

    # ── Policy documents ──────────────────────────────────────────────────────
    if not args.no_policies:
        print(f"Generating {len(ALL_POLICIES)} policy documents -> {POLICY_DIR}")
        for fn in ALL_POLICIES:
            fn()
    else:
        print("Skipping policy generation (--no-policies).")

    # ── Case documents ────────────────────────────────────────────────────────
    if not args.no_cases:
        cases = _select_cases(args)
        label = (
            f"{len(cases)} cases "
            f"(type={args.type}"
            + (f", count={args.count}" if args.count is not None else "")
            + f") -> {CASE_DIR}"
        )
        print(f"\nGenerating {label}")
        for case in cases:
            filename = f"{case['case_id'].lower()}.json"
            _save_case(case, filename)

        # Print summary table
        print(f"\n{'Case ID':<15} {'Type':<10} {'Expected':<25} {'LOB':<20} {'CPT'}")
        print("-" * 90)
        for c in cases:
            wf = c.get("workflow", "clinical")
            dec = c.get("expected_decision", "?")
            cpts = ", ".join(c.get("cpt_codes", []))
            print(f"{c['case_id']:<15} {wf[:9]:<10} {dec:<25} {c['lob']:<20} {cpts}")
    else:
        print("Skipping case generation (--no-cases).")

    print("\nDone.")


if __name__ == "__main__":
    main()
